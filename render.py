from docx import Document
from docx.shared import Pt, RGBColor
from datetime import date
import os

def save_brief(company_name, brief_text):
    doc = Document()

    # Title
    title = doc.add_heading(f'Pitch Brief: {company_name}', 0)
    title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Date
    doc.add_paragraph(f'Generated: {date.today().strftime("%d %B %Y")}')
    doc.add_paragraph('')

    # Body
    for line in brief_text.split('\n'):
        line = line.replace('**', '').replace('## ', '').replace('##', '')
        if line.strip() == '':
            doc.add_paragraph('')
        elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
            doc.add_heading(line.strip(), level=1)
        elif line.strip().startswith('-') or line.strip().startswith('•'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())

    # Save
    os.makedirs('output', exist_ok=True)
    filename = f"output/{company_name.replace(' ', '_')}_{date.today()}.docx"
    doc.save(filename)
    print(f"Brief saved to: {filename}")
    return filename

if __name__ == "__main__":
    from search import search_company
    from brief import generate_brief
    
    company = "Tinuiti"
    results = search_company(company)
    brief = generate_brief(company, results)
    save_brief(company, brief)